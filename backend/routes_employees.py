from fastapi import APIRouter, File, UploadFile, HTTPException, Request, Body, Depends
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
import tempfile
from mindee import Client, product, AsyncPredictResponse
import os
from db import get_connection
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from pypdf import PdfReader, PdfWriter
import shutil
import tempfile
from pypdf.generic import NameObject, BooleanObject, DictionaryObject
from docx import Document
import datetime
import pdfplumber
import re
from pdf_extract_utils import extract_einkommensbescheinigung_fields
from fastapi import Depends
from auth import get_current_user
import io
import datetime
import json

MINDEE_API_KEY = os.getenv("MINDEE_API_KEY", "your_mindee_api_key")
mindee_client = Client(api_key=MINDEE_API_KEY)

router = APIRouter()



@router.post("/employees/add")
async def add_employee(file: UploadFile = File(...)):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file.filename) as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save uploaded file: {str(e)}")
    try:
        input_doc = mindee_client.source_from_path(tmp_path)
        result: AsyncPredictResponse = mindee_client.enqueue_and_parse(
            product.InternationalIdV2,
            input_doc,
        )
        doc = result.document
        vorname = " ".join([n.value for n in getattr(doc, 'given_names', []) if hasattr(n, 'value') and n.value])
        # Try to extract geburtsname (birth name) if available
        geburtsname = None
        if hasattr(doc, 'surnames'):
            geburtsname = " ".join([s.value for s in getattr(doc, 'surnames', []) if hasattr(s, 'value') and s.value]) or None
        id_number = getattr(doc, 'document_number', None)
        id_number = id_number.value if id_number and hasattr(id_number, 'value') else None
        geburtsdatum = getattr(doc, 'birth_date', None)
        geburtsdatum = geburtsdatum.value if geburtsdatum and hasattr(geburtsdatum, 'value') else None
        geschlecht = getattr(doc, 'sex', None)
        geschlecht = geschlecht.value if geschlecht and hasattr(geschlecht, 'value') else None
        # Map M/F/D to German values
        if geschlecht == 'M':
            geschlecht = 'männlich'
        elif geschlecht == 'F':
            geschlecht = 'weiblich'
        elif geschlecht == 'D':
            geschlecht = 'divers'
        staatsangehoerigkeit = getattr(doc, 'nationality', None)
        staatsangehoerigkeit = staatsangehoerigkeit.value if staatsangehoerigkeit and hasattr(staatsangehoerigkeit, 'value') else None
        personal_number = getattr(doc, 'personal_number', None)
        personal_number = personal_number.value if personal_number and hasattr(personal_number, 'value') else None
        # Fallback extraction if Mindee fields are empty
        if not any([vorname, id_number, geburtsdatum, geschlecht, staatsangehoerigkeit, personal_number]):
            import re
            doc_str = str(doc)
            def extract(pattern):
                match = re.search(pattern, doc_str)
                return match.group(1).strip() if match else None
            id_number = extract(r"Document Number:\s*(.*)")
            vorname = extract(r"Given Names:\s*(.*)")
            geburtsname = extract(r"Surnames:\s*(.*)")
            geschlecht = extract(r"Sex:\s*(.*)")
            if geschlecht == 'M':
                geschlecht = 'männlich'
            elif geschlecht == 'F':
                geschlecht = 'weiblich'
            elif geschlecht == 'D':
                geschlecht = 'divers'
            geburtsdatum = extract(r"Birth Date:\s*(.*)")
            staatsangehoerigkeit = extract(r"Nationality:\s*(.*)")
            personal_number = extract(r"Personal Number:\s*(.*)")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Mindee extraction error: {str(e)}")
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        # Prevent duplicate id_number
        cursor.execute("SELECT id FROM employees WHERE id_number = %s", (id_number,))
        existing = cursor.fetchone()
        if existing:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=400, detail="Employee with this ID Number already exists.")
        cursor.execute(
            """
            INSERT INTO employees (
                vorname, geburtsname, geburtsdatum, geschlecht, staatsangehoerigkeit, id_number, personal_number
            ) VALUES (%s, %s, %s, %s, %s, %s, %s)
            """,
            (vorname, geburtsname, geburtsdatum, geschlecht, staatsangehoerigkeit, id_number, personal_number)
        )
        conn.commit()
        cursor.execute("SELECT * FROM employees WHERE id = LAST_INSERT_ID()")
        new_emp = cursor.fetchone()
        cursor.close()
        conn.close()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")
    return {"message": "Employee added successfully", "employee": new_emp}

@router.get("/employees/list")
def list_employees():
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM employees")
        employees = cursor.fetchall()
        cursor.close()
        conn.close()
        expected_fields = [
            "id", "id_number", "personal_number", "vorname", "geburtsname", "strasse_hausnummer", "plz_ort", "geburtsdatum", "geschlecht", "versicherungsnummer", "familienstand", "geburtsort_land", "schwerbehindert", "staatsangehoerigkeit", "arbeitnehmernummer", "iban", "bic", "eintrittsdatum", "ersteintrittsdatum", "betriebsstaette", "berufsbezeichnung", "taetigkeit", "hauptbeschaeftigung", "nebenbeschaeftigung", "weitere_beschaeftigungen", "schulabschluss", "berufsausbildung", "ausbildung_beginn", "ausbildung_ende", "baugewerbe_seit", "arbeitszeit_vollzeit", "arbeitszeit_teilzeit", "arbeitszeit_verteilung", "urlaubsanspruch", "kostenstelle", "abteilungsnummer", "personengruppe", "arbeitsverhaeltnis_befristet", "zweckbefristet", "befristung_arbeitsvertrag_zum", "schriftlicher_abschluss", "abschluss_arbeitsvertrag_am", "befristete_beschaeftigung_2monate", "weitere_angaben", "identifikationsnummer", "finanzamt_nr", "steuerklasse", "kinderfreibetraege", "konfession", "gesetzliche_krankenkasse", "elterneigenschaft", "kv", "rv", "av", "pv", "uv_gefahrtarif", "entlohnung_bezeichnung1", "entlohnung_betrag1", "entlohnung_gueltig_ab1", "entlohnung_stundenlohn1", "entlohnung_gueltig_ab_stunden1", "entlohnung_bezeichnung2", "entlohnung_betrag2", "entlohnung_gueltig_ab2", "entlohnung_stundenlohn2", "entlohnung_gueltig_ab_stunden2", "entlohnung_bezeichnung3", "entlohnung_betrag3", "entlohnung_gueltig_ab3", "entlohnung_stundenlohn3", "entlohnung_gueltig_ab_stunden3", "vwl_empfaenger", "vwl_betrag", "vwl_ag_anteil", "vwl_seit_wann", "vwl_vertragsnr", "vwl_kontonummer", "vwl_bankleitzahl", "ap_arbeitsvertrag", "ap_bescheinigung_lsta", "ap_sv_ausweis", "ap_mitgliedsbescheinigung_kk", "ap_bescheinigung_private_kk", "ap_vwl_vertrag", "ap_nachweis_elterneigenschaft", "ap_vertrag_bav", "ap_schwerbehindertenausweis", "ap_unterlagen_sozialkasse", "vorbeschaeftigung_zeitraum_von", "vorbeschaeftigung_zeitraum_bis", "vorbeschaeftigung_art", "vorbeschaeftigung_tage"
        ]
        for emp in employees:
            for field in expected_fields:
                v = emp.get(field, '')
                if v is None:
                    emp[field] = ''
                elif not isinstance(v, str):
                    emp[field] = str(v)
        return JSONResponse(content=employees, status_code=200)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

# All endpoints below this require authentication
@router.delete("/employees/delete/{employee_id}")
def delete_employee(employee_id: int, user=Depends(get_current_user)):
    try:
        conn = get_connection()
        cursor = conn.cursor()
        
        # First, delete related records from erklaerung_form table
        cursor.execute("DELETE FROM erklaerung_form WHERE employee_id = %s", (employee_id,))
        
        # Then delete related records from einkommensbescheinigung table if it exists
        cursor.execute("DELETE FROM einkommensbescheinigung WHERE employee_id = %s", (employee_id,))
        
        # Finally delete the employee
        cursor.execute("DELETE FROM employees WHERE id = %s", (employee_id,))
        
        conn.commit()
        cursor.close()
        conn.close()
        return {"message": f"Employee with id {employee_id} deleted successfully."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

@router.patch("/employees/edit/{employee_id}")
async def edit_employee(employee_id: int, req: Request, user=Depends(get_current_user)):
    try:
        data = await req.json()
    except Exception:
        data = {}
    if not data:
        raise HTTPException(status_code=400, detail="No data provided for update.")
    # Map M/F/D to German values if present in edit
    if 'geschlecht' in data:
        if data['geschlecht'] == 'M':
            data['geschlecht'] = 'männlich'
        elif data['geschlecht'] == 'F':
            data['geschlecht'] = 'weiblich'
        elif data['geschlecht'] == 'D':
            data['geschlecht'] = 'divers'
    # Only allow updating fields that exist in the new schema
    allowed_fields = [
        "vorname", "geburtsname", "strasse_hausnummer", "plz_ort", "geburtsdatum", "geschlecht", "versicherungsnummer", "familienstand", "geburtsort_land", "schwerbehindert", "staatsangehoerigkeit", "arbeitnehmernummer", "iban", "bic", "eintrittsdatum", "ersteintrittsdatum", "betriebsstaette", "berufsbezeichnung", "taetigkeit", "hauptbeschaeftigung", "nebenbeschaeftigung", "weitere_beschaeftigungen", "schulabschluss", "berufsausbildung", "ausbildung_beginn", "ausbildung_ende", "baugewerbe_seit", "arbeitszeit_vollzeit", "arbeitszeit_teilzeit", "arbeitszeit_verteilung", "urlaubsanspruch", "kostenstelle", "abteilungsnummer", "personengruppe", "arbeitsverhaeltnis_befristet", "zweckbefristet", "befristung_arbeitsvertrag_zum", "schriftlicher_abschluss", "abschluss_arbeitsvertrag_am", "befristete_beschaeftigung_2monate", "weitere_angaben", "identifikationsnummer", "finanzamt_nr", "steuerklasse", "kinderfreibetraege", "konfession", "gesetzliche_krankenkasse", "elterneigenschaft", "kv", "rv", "av", "pv", "uv_gefahrtarif", "entlohnung_bezeichnung1", "entlohnung_betrag1", "entlohnung_gueltig_ab1", "entlohnung_stundenlohn1", "entlohnung_gueltig_ab_stunden1", "entlohnung_bezeichnung2", "entlohnung_betrag2", "entlohnung_gueltig_ab2", "entlohnung_stundenlohn2", "entlohnung_gueltig_ab_stunden2", "entlohnung_bezeichnung3", "entlohnung_betrag3", "entlohnung_gueltig_ab3", "entlohnung_stundenlohn3", "entlohnung_gueltig_ab_stunden3", "vwl_empfaenger", "vwl_betrag", "vwl_ag_anteil", "vwl_seit_wann", "vwl_vertragsnr", "vwl_kontonummer", "vwl_bankleitzahl", "ap_arbeitsvertrag", "ap_bescheinigung_lsta", "ap_sv_ausweis", "ap_mitgliedsbescheinigung_kk", "ap_bescheinigung_private_kk", "ap_vwl_vertrag", "ap_nachweis_elterneigenschaft", "ap_vertrag_bav", "ap_schwerbehindertenausweis", "ap_unterlagen_sozialkasse", "vorbeschaeftigung_zeitraum_von", "vorbeschaeftigung_zeitraum_bis", "vorbeschaeftigung_art", "vorbeschaeftigung_tage", "id_number", "personal_number"
    ]
    fields = []
    values = []
    for field in allowed_fields:
        if field in data:
            # If value is None or empty string, store as NULL in DB
            val = data[field]
            if val is None or (isinstance(val, str) and val.strip() == ""):
                val = None
            fields.append(f"{field} = %s")
            values.append(val)
    if not fields:
        raise HTTPException(status_code=400, detail="No valid fields provided for update.")
    values.append(employee_id)
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute(f"UPDATE employees SET {', '.join(fields)} WHERE id = %s", tuple(values))
        conn.commit()
        cursor.execute("SELECT * FROM employees WHERE id = %s", (employee_id,))
        updated = cursor.fetchone()
        cursor.close()
        conn.close()
        if not updated:
            raise HTTPException(status_code=404, detail="Employee not found.")
        return {"message": "Employee updated successfully", "employee": updated}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

@router.get("/employees/pdf/{employee_id}")
def download_employee_pdf(employee_id: int, user=Depends(get_current_user)):
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM employees WHERE id = %s", (employee_id,))
        emp = cursor.fetchone()
        cursor.close()
        conn.close()
        if not emp:
            raise HTTPException(status_code=404, detail="Employee not found.")
        return emp
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation error: {str(e)}")

@router.get("/arbeitsvertrag/list")
def arbeitsvertrag_list(user=Depends(get_current_user)):
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute('''
            SELECT
                e.id AS id,
                CONCAT(e.vorname, ' ', e.geburtsname) AS name,
                e.strasse_hausnummer AS strasse,
                e.plz_ort AS plz_ort,
                e.land AS land,
                e.contract_type AS contract_type,
                f.beschaeftigung_beginn AS beginn,
                f.beschaeftigung_berufsbezeichnung AS position,
                f.arbeitszeit_stunden AS arbeitszeit_stunden,
                f.entgelt_pro_monat_wert AS gehalt,
                f.urlaubsanspruch_tage AS urlaub
            FROM employees e
            JOIN erklaerung_form f ON e.id = f.employee_id
        ''')
        rows = cursor.fetchall()
        # Replace None with '' for frontend compatibility
        for row in rows:
            for k, v in row.items():
                if v is None:
                    row[k] = ''
        cursor.close()
        conn.close()
        return rows
    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

@router.patch("/arbeitsvertrag/edit/{employee_id}")
def arbeitsvertrag_edit(employee_id: int, data: dict = Body(...), user=Depends(get_current_user)):
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        # Employees table fields
        emp_fields = {}
        if 'name' in data:
            parts = data['name'].split(' ', 1)
            emp_fields['vorname'] = parts[0]
            emp_fields['geburtsname'] = parts[1] if len(parts) > 1 else ''
        for f in ['strasse', 'plz_ort', 'land', 'contract_type']:
            if f in data:
                dbf = 'strasse_hausnummer' if f == 'strasse' else f
                emp_fields[dbf] = data[f]
        if emp_fields:
            sets = ', '.join([f"{k} = %s" for k in emp_fields])
            vals = list(emp_fields.values()) + [employee_id]
            cursor.execute(f"UPDATE employees SET {sets} WHERE id = %s", vals)
        # Erklaerung_form table fields
        erk_fields = {}
        if 'beginn' in data:
            erk_fields['beschaeftigung_beginn'] = data['beginn']
        if 'position' in data:
            erk_fields['beschaeftigung_berufsbezeichnung'] = data['position']
        if 'arbeitszeit_stunden' in data:
            erk_fields['arbeitszeit_stunden'] = data['arbeitszeit_stunden']
        if 'gehalt' in data:
            erk_fields['entgelt_pro_monat_wert'] = data['gehalt']
        if 'urlaub' in data:
            erk_fields['urlaubsanspruch_tage'] = data['urlaub']
        if erk_fields:
            sets = ', '.join([f"{k} = %s" for k in erk_fields])
            vals = list(erk_fields.values()) + [employee_id]
            cursor.execute(f"UPDATE erklaerung_form SET {sets} WHERE employee_id = %s", vals)
        conn.commit()
        # Return updated row
        cursor.execute('''
            SELECT
                e.id AS id,
                CONCAT(e.vorname, ' ', e.geburtsname) AS name,
                e.strasse_hausnummer AS strasse,
                e.plz_ort AS plz_ort,
                e.land AS land,
                e.contract_type AS contract_type,
                f.beschaeftigung_beginn AS beginn,
                f.beschaeftigung_berufsbezeichnung AS position,
                f.arbeitszeit_stunden AS arbeitszeit_stunden,
                f.entgelt_pro_monat_wert AS gehalt,
                f.urlaubsanspruch_tage AS urlaub
            FROM employees e
            JOIN erklaerung_form f ON e.id = f.employee_id
            WHERE e.id = %s
        ''', (employee_id,))
        row = cursor.fetchone()
        for k, v in row.items():
            if v is None:
                row[k] = ''
        cursor.close()
        conn.close()
        return row
    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

@router.get("/arbeitsvertrag/download/{employee_id}")
def arbeitsvertrag_download(employee_id: int):
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute('''
            SELECT
                e.id AS id,
                CONCAT(e.vorname, ' ', e.geburtsname) AS name,
                e.strasse_hausnummer AS strasse,
                e.plz_ort AS plz_ort,
                e.land AS land,
                e.contract_type AS contract_type,
                f.beschaeftigung_beginn AS beginn,
                f.beschaeftigung_berufsbezeichnung AS position,
                f.arbeitszeit_stunden AS arbeitszeit_stunden,
                f.entgelt_pro_monat_wert AS gehalt,
                f.urlaubsanspruch_tage AS urlaub
            FROM employees e
            JOIN erklaerung_form f ON e.id = f.employee_id
            WHERE e.id = %s
        ''', (employee_id,))
        row = cursor.fetchone()
        cursor.close()
        conn.close()
        if not row:
            raise HTTPException(status_code=404, detail="Employee/contract not found.")
        # Allow download for all contract types and select template accordingly
        contract_type = row['contract_type']
        if not contract_type:
            raise HTTPException(status_code=400, detail="Vertragsart nicht gesetzt.")
        # Map contract type to template file
        template_map = {
            'TEILZEITTÄTIGKEIT': 'Arbeitsvertrag Teilzeit.docx',
            'VOLLZEITTÄTIGKEIT': 'Arbeitsvertrag fuer Arbeitnehmer - Vollzeit .docx',
            'TEILZEITTÄTIGKEIT - "MINIJOB"': 'Minijob-Vertrag.docx',
        }
        template_file = template_map.get(contract_type)
        if not template_file:
            raise HTTPException(status_code=400, detail=f"Kein Template für Vertragsart: {contract_type}")
        template_path = os.path.join(os.path.dirname(__file__), 'filesDoc', template_file)
        if not os.path.exists(template_path):
            raise HTTPException(status_code=500, detail=f"Template nicht gefunden: {template_file}")
        doc = Document(template_path)
        # Replace placeholders
        def to_str(val):
            if isinstance(val, (datetime.date, datetime.datetime)):
                return val.strftime("%d.%m.%Y")
            return str(val) if val is not None else ''
        replacements = {
            '{{name}}': to_str(row['name']),
            '{{strasse}}': to_str(row['strasse']),
            '{{plz_ort}}': to_str(row['plz_ort']),
            '{{land}}': to_str(row['land']),
            '{{beginn}}': to_str(row['beginn']),
            '{{position}}': to_str(row['position']),
            '{{arbeitszeit_stunden}}': to_str(row.get('arbeitszeit_stunden', '')),
            '{{gehalt}}': to_str(row['gehalt']),
            '{{urlaub}}': to_str(row['urlaub']),
        }
        # Replace in paragraphs (handles split runs)
        for p in doc.paragraphs:
            for k, v in replacements.items():
                if k in p.text:
                    p.text = p.text.replace(k, v)
        # Replace in tables
        for table in doc.tables:
            for row_cells in table.rows:
                for cell in row_cells.cells:
                    for k, v in replacements.items():
                        if k in cell.text:
                            cell.text = cell.text.replace(k, v)
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        # Return file
        filename = f"Arbeitsvertrag_{contract_type}_{row['name'].replace(' ', '_')}.docx"
        return FileResponse(tmp_path, filename=filename, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={"Content-Disposition": f"attachment; filename={filename}"})
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print('--- Arbeitsvertrag Download Error ---')
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Docx generation error: {str(e)}")

@router.get("/employees/stundenzettel-data/{employee_id}")
def get_stundenzettel_data(employee_id: int, year: int, user=Depends(get_current_user)):
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM employees WHERE id = %s", (employee_id,))
        emp = cursor.fetchone()
        if not emp:
            cursor.close()
            conn.close()
            raise HTTPException(status_code=404, detail="Employee not found.")
        cursor.execute("SELECT * FROM company LIMIT 1")
        company = cursor.fetchone()
        # Fetch daily entries for the year
        cursor.execute("""
            SELECT * FROM einkommensbescheinigung
            WHERE employee_id = %s AND jahr = %s
        """, (employee_id, str(year)))
        rows = cursor.fetchall()
        # Build entries dict: entries[month][day] = {...}
        entries = {}
        for row in rows:
            month = int(row.get('monat', '0'))
            # Assume 'tag' field exists, else skip (if not, you need to add day info to your table)
            day = row.get('tag')
            if not day:
                continue
            day = str(day).zfill(2)
            if month not in entries:
                entries[month] = {}
            entries[month][day] = {
                'beginn': row.get('beginn', ''),
                'pause': row.get('pause', ''),
                'ende': row.get('ende', ''),
                'dauer': row.get('dauer', ''),
                'code': row.get('code', ''),
                'aufgezeichnet_am': row.get('aufgezeichnet_am', ''),
                'bemerkungen': row.get('bemerkungen', ''),
            }
        cursor.close()
        conn.close()
        company_name = company["name"] if company else ""
        employee_name = f"{emp.get('vorname', '')} {emp.get('geburtsname', '')}".strip()
        employee_number = emp.get('personal_number', '')
        arbeitszeit_verteilung = emp.get('arbeitszeit_verteilung', '')
        return {
            "companyName": company_name,
            "employeeName": employee_name,
            "employeeNumber": employee_number,
            "year": year,
            "entries": entries,
            "arbeitszeitVerteilung": arbeitszeit_verteilung
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

# --- NEW: Stundenzettel Multi-Employee PDF and History ---

@router.post("/employees/stundenzettel-pdf")
def generate_stundenzettel_pdf(data: dict = Body(...), user=Depends(get_current_user)):
    employee_ids = data.get('employee_ids', [])
    month = int(data.get('month'))
    year = int(data.get('year'))
    if not employee_ids or not month or not year:
        raise HTTPException(status_code=400, detail="Missing parameters.")
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    employees = []
    for eid in employee_ids:
        cursor.execute("SELECT * FROM employees WHERE id = %s", (eid,))
        emp = cursor.fetchone()
        if not emp:
            continue
        cursor.execute("SELECT * FROM company LIMIT 1")
        company = cursor.fetchone()
        cursor.execute("SELECT * FROM einkommensbescheinigung WHERE employee_id = %s AND jahr = %s AND monat = %s", (eid, str(year), str(month)))
        rows = cursor.fetchall()
        entries = {}
        for row in rows:
            day = row.get('tag')
            if not day:
                continue
            day = str(day).zfill(2)
            entries[day] = {
                'beginn': row.get('beginn', ''),
                'pause': row.get('pause', ''),
                'ende': row.get('ende', ''),
                'dauer': row.get('dauer', ''),
                'code': row.get('code', ''),
                'aufgezeichnet_am': row.get('aufgezeichnet_am', ''),
                'bemerkungen': row.get('bemerkungen', ''),
            }
        employees.append({
            'companyName': company["name"] if company else "",
            'employeeName': f"{emp.get('vorname', '')} {emp.get('geburtsname', '')}".strip(),
            'employeeNumber': emp.get('personal_number', ''),
            'entries': entries,
            'arbeitszeitVerteilung': emp.get('arbeitszeit_verteilung', ''),
        })
    cursor.close()
    conn.close()
    if not employees:
        raise HTTPException(status_code=404, detail="No employees found.")
    # Generate PDF using ReportLab (placeholder)
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    import io
    import json
    from datetime import datetime
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    c.drawString(100, 800, f"Stundenzettel PDF for {len(employees)} employees, {month}/{year}")
    for idx, emp in enumerate(employees):
        c.drawString(100, 780 - idx*20, f"{emp['employeeName']} ({emp['employeeNumber']})")
    c.showPage()
    c.save()
    buffer.seek(0)
    # Save download log
    conn = get_connection()
    cursor = conn.cursor()
    employee_names = [emp['employeeName'] for emp in employees]
    filename = f"Stundenzettel_{year}_{month}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
    file_blob = buffer.getvalue()
    cursor.execute(
        """
        INSERT INTO stundenzettel_downloads (user_id, employee_ids, employee_names, month, year, download_date, file_blob, filename)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """,
        (
            user['id'],
            json.dumps(employee_ids),
            json.dumps(employee_names),
            month,
            year,
            datetime.now(),
            file_blob,
            filename
        )
    )
    conn.commit()
    cursor.close()
    conn.close()
    return StreamingResponse(io.BytesIO(file_blob), media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={filename}"})

@router.post("/employees/stundenzettel-pdf-data")
def get_stundenzettel_pdf_data(data: dict = Body(...), user=Depends(get_current_user)):
    employee_ids = data.get('employee_ids', [])
    month = int(data.get('month'))
    year = int(data.get('year'))
    if not employee_ids or not month or not year:
        raise HTTPException(status_code=400, detail="Missing parameters.")
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    employees = []
    from calendar import monthrange
    import math
    def hours_to_hhmm(hours):
        if not hours or math.isnan(hours):
            return ''
        h = int(hours)
        m = int(round((hours - h) * 60))
        return f"{h:02d}:{m:02d}"
    weekday_map = ['Mo', 'Di', 'Mi', 'Do', 'Fr', 'Sa', 'So']
    for eid in employee_ids:
        cursor.execute("SELECT * FROM employees WHERE id = %s", (eid,))
        emp = cursor.fetchone()
        if not emp:
            continue
        cursor.execute("SELECT * FROM company LIMIT 1")
        company = cursor.fetchone()
        # Get arbeitszeit_verteilung as dict {weekday: hours}
        arbeitszeit_verteilung = emp.get('arbeitszeit_verteilung', '')
        verteilung = {}
        for part in (arbeitszeit_verteilung or '').split(','):
            if ':' in part:
                w, h = part.split(':')
                try:
                    verteilung[w.strip()] = float(h.strip())
                except:
                    pass
        # Fetch all entries for the month
        cursor.execute("SELECT * FROM einkommensbescheinigung WHERE employee_id = %s AND jahr = %s AND monat = %s", (eid, str(year), str(month)))
        rows = cursor.fetchall()
        # Build entries dict for all days in month
        entries = {}
        num_days = monthrange(year, month)[1]
        for day in range(1, num_days + 1):
            day_str = str(day).zfill(2)
            # Find DB entry for this day
            row = next((r for r in rows if str(r.get('tag')).zfill(2) == day_str), None)
            if row:
                dauer = row.get('dauer', '')
            else:
                # Get weekday for this date
                import datetime
                dt = datetime.date(year, month, day)
                weekday = weekday_map[dt.weekday()]  # 0=Monday
                hours = verteilung.get(weekday, 0)
                dauer = hours_to_hhmm(hours) if hours else ''
            entries[day_str] = {
                'beginn': row.get('beginn', '') if row else '',
                'pause': row.get('pause', '') if row else '',
                'ende': row.get('ende', '') if row else '',
                'dauer': dauer,
                'code': row.get('code', '') if row else '',
                'aufgezeichnet_am': row.get('aufgezeichnet_am', '') if row else '',
                'bemerkungen': row.get('bemerkungen', '') if row else '',
            }
        employees.append({
            'companyName': company["name"] if company else "",
            'employeeName': f"{emp.get('vorname', '')} {emp.get('geburtsname', '')}".strip(),
            'employeeNumber': emp.get('personal_number', ''),
            'entries': entries,
            'arbeitszeitVerteilung': arbeitszeit_verteilung,
        })
    cursor.close()
    conn.close()
    if not employees:
        raise HTTPException(status_code=404, detail="No employees found.")
    return {"employees": employees, "month": month, "year": year}

@router.post("/employees/stundenzettel-log-download")
def log_stundenzettel_download(
    employee_ids: str = Body(...),
    month: int = Body(...),
    year: int = Body(...),
    employee_names: str = Body(...),
    file: UploadFile = File(None),
    user=Depends(get_current_user)
):
    import json as pyjson
    # Parse JSON fields if sent as strings
    if isinstance(employee_ids, str):
        employee_ids = pyjson.loads(employee_ids)
    if isinstance(employee_names, str):
        employee_names = pyjson.loads(employee_names)
    filename = f"Stundenzettel_{year}_{month}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
    file_blob = None
    if file is not None:
        file_blob = file.file.read()
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute(
        """
        INSERT INTO stundenzettel_downloads (user_id, employee_ids, employee_names, month, year, download_date, filename, file_blob)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """,
        (
            user['id'],
            pyjson.dumps(employee_ids),
            pyjson.dumps(employee_names),
            month,
            year,
            datetime.now(),
            filename,
            file_blob
        )
    )
    conn.commit()
    cursor.close()
    conn.close()
    return {"message": "Download logged with file"}

@router.get("/employees/stundenzettel-history")
def get_stundenzettel_history(user=Depends(get_current_user)):
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT id, employee_names, month, year, download_date, filename FROM stundenzettel_downloads WHERE user_id = %s ORDER BY download_date DESC", (user['id'],))
    rows = cursor.fetchall()
    for row in rows:
        try:
            row['employee_names'] = json.loads(row['employee_names'])
        except:
            row['employee_names'] = []
    cursor.close()
    conn.close()
    return rows

@router.get("/employees/stundenzettel-history/{download_id}/download")
def download_stundenzettel_history(download_id: int, user=Depends(get_current_user)):
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT file_blob, filename, user_id FROM stundenzettel_downloads WHERE id = %s", (download_id,))
    row = cursor.fetchone()
    cursor.close()
    conn.close()
    if not row or row['user_id'] != user['id']:
        raise HTTPException(status_code=404, detail="Not found.")
    return StreamingResponse(io.BytesIO(row['file_blob']), media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename={row['filename']}"})